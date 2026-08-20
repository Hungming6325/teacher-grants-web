from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def apply_font(run, ascii_name="Calibri", east_asia="Microsoft JhengHei"):
    run.font.name = ascii_name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, MID_GRAY, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def base_doc(running_title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    hp = section.header.paragraphs[0]
    hp.text = running_title
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        apply_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    fp = section.footer.paragraphs[0]
    add_page_number(fp)
    for run in fp.runs:
        apply_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    return doc


def cover(doc, kicker, title, subtitle, audience):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    apply_font(r)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    apply_font(r)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    apply_font(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"適用對象：{audience}")
    apply_font(r)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"文件版本：1.0　｜　產出日期：{date.today().isoformat()}")
    apply_font(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    doc.add_page_break()


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        apply_font(r)
    return p


def add_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        apply_font(r)
        r.bold = True
        r = p.add_run(text[len(bold_prefix):])
        apply_font(r)
    else:
        r = p.add_run(text)
        apply_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        apply_font(r)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        apply_font(r)


def add_note(doc, label, text, tone="blue"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF4FB" if tone == "blue" else "FFF7E6")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE if tone == "blue" else GOLD)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}：")
    apply_font(r)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE if tone == "blue" else GOLD)
    r = p.add_run(text)
    apply_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        apply_font(r)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            apply_font(r)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_architecture():
    doc = base_doc("教師教研數據管理平台｜系統架構說明")
    cover(doc, "System Architecture", "教師教研數據管理平台\n系統架構說明", "依 GitHub clone 專案原始碼盤點之現況架構", "系統維運人員、開發人員、資料管理人員")
    add_h(doc, "1. 文件目的與範圍")
    add_p(doc, "本文件依目前儲存庫內容說明系統組成、資料流、模組責任、部署方式與維運注意事項。內容反映 2026-08-17 的程式碼狀態，不將尚未實作的登入、資料庫、API 或後台上傳功能描述為既有功能。")
    add_note(doc, "架構結論", "本系統是 Next.js 16 App Router 的唯讀分析網站。伺服器元件於請求時從專案 data 目錄讀取 CSV，解析後將資料傳入瀏覽器端儀表板；使用者的篩選與圖表互動只發生在前端記憶體，不會回寫來源檔。")

    add_h(doc, "2. 技術堆疊")
    add_table(doc, ["層級", "技術／套件", "用途"], [
        ("框架", "Next.js 16.2.1、React 19.2.4、TypeScript 5", "App Router、伺服器／用戶端元件與型別安全"),
        ("介面", "Tailwind CSS 4、Framer Motion、Lucide React", "響應式版面、動態效果與圖示"),
        ("圖表", "Recharts 3.8.1、react-countup", "圓餅、折線、長條、樹狀與數字動畫"),
        ("資料", "Papa Parse、Node.js fs/path", "CSV 解析與伺服器端檔案讀取"),
        ("其他", "xlsx", "已列入依賴；目前核心頁面未發現使用者匯出流程"),
    ], [1700, 3000, 4660])

    add_h(doc, "3. 邏輯架構與資料流")
    add_steps(doc, [
        "瀏覽器請求首頁或五個儀表板路由。",
        "app/<route>/page.tsx 伺服器元件呼叫對應的 *-server.ts 載入器。",
        "載入器透過 src/lib/data-files.ts 統一解析 process.cwd()/data/<file>.csv 路徑，並以 UTF-8 讀取。",
        "領域解析器使用 Papa Parse 將 CSV 欄位清理、正規化、轉型與過濾，產生 TypeScript record。",
        "伺服器元件把 record 陣列序列化為 client dashboard 的 props。",
        "瀏覽器使用 React state/useMemo 執行篩選、彙總及圖表互動；不呼叫寫入 API，也不修改 CSV。",
    ])
    add_p(doc, "資料流簡圖：瀏覽器 → App Router 頁面 → Server loader → data/*.csv → Parser/Normalizer → Client Dashboard → 篩選、統計與圖表。")

    add_h(doc, "4. 路由與功能模組")
    add_table(doc, ["路由", "伺服器資料來源", "用戶端儀表板", "主要內容"], [
        ("/", "無", "入口頁", "五大分析模組導覽"),
        ("/grants", "grants_112_114.csv", "TeacherGrantsDashboard", "獎補助金額、教師、年度占比與次項目趨勢"),
        ("/teaching", "teaching_awards_114_115.csv", "TeachingAnalyticsDashboard", "教學獎勵點數、金額、申請與合作關係"),
        ("/papers", "papers_112_114.csv", "PaperPublicationsDashboard", "論文去重、收錄分類、作者角色、合著與趨勢"),
        ("/projects", "projects_112_114.csv", "ProjectContractsDashboard", "計畫件數、金額、類型、系所／教師排行與清單"),
        ("/teachers", "teachers、papers、projects、patent、transfer", "FacultyStaffingDashboard", "教師人力結構與研究評鑑風險"),
    ], [1200, 2200, 2600, 3360])

    add_h(doc, "5. 程式目錄責任")
    add_table(doc, ["位置", "責任"], [
        ("app/", "App Router 路由、全域 layout、Metadata 與伺服器端資料載入入口。"),
        ("src/components/dashboard/", "五個互動式儀表板；持有篩選狀態、彙總呈現與圖表選取狀態。"),
        ("src/components/charts/", "可重用圖表元件。"),
        ("src/components/ui/", "PanelCard、DashboardTabs 等共用介面。"),
        ("src/lib/*-server.ts", "僅伺服器使用的 CSV 檔案讀取。"),
        ("src/lib/*.ts", "CSV 解析、欄位清理、篩選、去重、統計、個資遮罩與資料檔映射。"),
        ("src/types/", "各領域 record、filter、summary 型別。"),
        ("data/", "七份原始 CSV；目前等同系統資料儲存層。"),
        ("scripts/", "資料正規化與報表／文件輔助腳本。"),
    ], [2500, 6860])

    add_h(doc, "6. 資料來源與目前快照")
    add_table(doc, ["檔案", "原始資料列數", "用途"], [
        ("grants_112_114.csv", "5,753", "獎補助"),
        ("teaching_awards_114_115.csv", "3,890", "教學精進"),
        ("papers_112_114.csv", "1,193", "論文與教師風險"),
        ("projects_112_114.csv", "1,382", "計畫與教師風險"),
        ("teachers_114_2.csv", "415", "教師主檔"),
        ("patent.csv", "383", "教師研究成果／風險"),
        ("transfer.csv", "12", "教師研究成果／風險"),
    ], [3800, 1900, 3660])
    add_note(doc, "注意", "以上為儲存庫目前 CSV 的實體列數，不等同畫面統計值。解析器會排除缺少必要欄位的資料，論文頁還會依發表年份與正規化論文名稱去重。", "gold")

    add_h(doc, "7. 關鍵領域規則")
    add_h(doc, "7.1 個資顯示", 2)
    add_p(doc, "獎補助與教學精進頁透過 src/lib/privacy.ts 遮罩教師姓名：兩字姓名保留首字加 O；三字以上保留首尾字，中間以 O 取代。教師聘任與研究風險頁目前會顯示完整姓名，部署前應依組織權限政策再次確認。")
    add_h(doc, "7.2 論文去重與分類", 2)
    add_bullets(doc, [
        "同一論文以「發表年份 + 正規化論文名稱」建立主鍵；沒有標題時改用年份、期刊與教師。",
        "SCI 正規化為 SCIE；SSCI only、SCIE only、SSCI + SCIE 為互斥桶。",
        "同一篇論文可彙整多位教師、系所、收錄分類、第一作者與通訊作者。",
        "校內合著以同一去重論文含兩位以上校內教師判定。",
    ])
    add_h(doc, "7.3 教師研究評鑑風險", 2)
    add_bullets(doc, [
        "候選人：專任、講師以上、服務滿 3 年且年齡小於 62 歲。",
        "缺少計畫：計畫件數為 0。",
        "缺少成果：期刊論文、專利與技術移轉合計為 0。",
        "高風險：同時缺少計畫與成果；中風險：僅缺少其中一項。",
        "此為系統篩選規則，不等同正式人事決策；資料期間與姓名匹配品質會直接影響結果。",
    ])

    add_h(doc, "8. 執行與部署")
    add_h(doc, "8.1 本機開發", 2)
    add_steps(doc, [
        "安裝 Node.js 與 npm，於專案根目錄執行 npm install。",
        "執行 npm run dev；目前 script 使用 next dev --webpack。",
        "以瀏覽器開啟 http://localhost:3000。",
        "修改程式後執行 npm run lint 與 npm run build 驗證。",
    ])
    add_h(doc, "8.2 正式環境", 2)
    add_p(doc, "標準流程為 npm run build 後以 npm run start 啟動。部署套件必須包含 data/ 目錄，且執行帳號須能讀取 CSV。若平台採不可變映像，更新 CSV 後必須重新建置／部署；若採掛載磁碟，則需設計一致的檔案更新、版本與回復流程。")

    add_h(doc, "9. CSV 更新維運程序")
    add_steps(doc, [
        "先備份既有 CSV，並確認新檔為 UTF-8、首列欄名與既有格式一致。",
        "若檔名變更，同步修改 src/lib/data-files.ts；不要只改 data/README.md。",
        "抽查必要欄位、數字格式、年度、系所與教師姓名是否可被解析。",
        "執行 npm run lint 與 npm run build。",
        "逐頁驗證預設篩選、總數、空資料狀態、圖表點選與重設功能。",
        "提交 CSV、映射與文件變更至 Git，保留來源日期與版本紀錄。",
    ])

    add_h(doc, "10. 安全、隱私與限制")
    add_bullets(doc, [
        "目前未見登入、角色權限、資料列級授權或稽核軌跡；不應直接暴露於未受控公開網路。",
        "CSV 可能含姓名、電子郵件、年齡、學歷與聘任資料；儲存庫與部署環境皆應採最小權限。",
        "瀏覽器收到儀表板 props 後可能包含比畫面可見更多的欄位；姓名遮罩不等同資料存取控制。",
        "系統沒有寫入或後台管理功能；更新資料需由維運人員替換檔案並重新部署或重啟。",
        "大量 CSV 會增加伺服器讀取、序列化與用戶端記憶體成本；成長後宜改為資料庫／查詢 API 與伺服器端聚合。",
        "錯誤處理以 console.error 為主，缺少集中監控、結構化日誌與資料品質告警。",
    ])

    add_h(doc, "11. 建議演進路線")
    add_table(doc, ["優先級", "建議", "目的"], [
        ("P0", "部署前加入身份驗證、角色授權與受保護網路邊界", "保護教師個資與風險名單"),
        ("P0", "建立 CSV schema 驗證、更新檢核與錯誤頁", "防止欄位漂移造成錯誤統計或整頁失敗"),
        ("P1", "建立資料版本、來源日期與最後更新時間", "提升可追溯性"),
        ("P1", "加入單元測試覆蓋去重、金額、風險與篩選規則", "降低規則回歸"),
        ("P2", "資料量擴大時導入資料庫、API、快取與伺服器聚合", "改善效能與治理"),
        ("P2", "導入集中日誌、錯誤監控與健康檢查", "縮短維運問題定位時間"),
    ], [1300, 4900, 3160])

    add_h(doc, "附錄 A｜關鍵設定與檔案")
    add_bullets(doc, [
        "package.json：依賴與 dev/build/start/lint 指令。",
        "next.config.ts：Next.js 設定。",
        "app/layout.tsx：zh-Hant 根版面與網站 Metadata。",
        "src/lib/data-files.ts：CSV 檔名的單一映射來源。",
        "data/README.md：資料檔與頁面對照。",
        "src/lib/privacy.ts：姓名遮罩。",
        "src/lib/papers.ts：論文正規化與去重。",
        "src/components/dashboard/FacultyStaffingDashboard.tsx：研究風險規則整合。",
    ])
    path = OUT / "系統架構說明.docx"
    doc.save(path)
    return path


def build_manual():
    doc = base_doc("教師教研數據管理平台｜使用者操作手冊")
    cover(doc, "User Guide", "教師教研數據管理平台\n使用者操作手冊", "五大儀表板的篩選、圖表互動與判讀方式", "校內主管、行政人員、研究與教學業務承辦人")
    add_h(doc, "1. 系統用途")
    add_p(doc, "本平台整合教師獎補助、教學精進、論文發表、計畫承接及教師聘任資料，提供互動式統計與視覺化。系統目前為查詢／分析用途，畫面操作不會修改原始 CSV。")
    add_note(doc, "使用前提醒", "畫面內容依系統目前載入的資料期間與品質計算。若要作正式決策或對外報送，請先確認資料版本、篩選條件與業務定義。")

    add_h(doc, "2. 進入系統與共通導覽")
    add_steps(doc, [
        "使用瀏覽器開啟系統網址；本機環境預設為 http://localhost:3000。",
        "首頁顯示五張入口卡：獎補助分析、教學精進、論文發表、計畫承接、教師聘任。",
        "點選卡片進入模組；儀表板右上方頁籤可直接切換其他模組或返回入口。",
        "在窄螢幕上向下捲動查看所有圖表；建議桌機使用 1280px 以上寬度。",
    ])
    add_h(doc, "2.1 共通操作觀念", 2)
    add_bullets(doc, [
        "篩選條件會立即套用，不需另外按查詢。",
        "多數圖表可點選資料項目作為交叉篩選；再次點選、按「清除選項」或「重設篩選」可還原。",
        "教師下拉選單支援關鍵字搜尋；部分模組會以 O 遮罩姓名。",
        "收錄分類、計畫類型、職級等多選欄位可同時勾選多項；「清除」會移除該欄位的選取。",
        "詳細清單預設收合，需按「顯示…」展開，以避免初始畫面過長。",
    ])

    add_h(doc, "3. 獎補助分析")
    add_p(doc, "用途：掌握 112–114 年教師獎補助投入、次項目分布與年度變化。預設年度為 114。")
    add_h(doc, "3.1 篩選", 2)
    add_bullets(doc, [
        "年度：112、113、114 或全部。",
        "系所：單選；教師名單會隨相關條件更新。",
        "教師：開啟下拉後輸入姓名關鍵字；畫面以遮罩姓名顯示。",
        "次項目：單選；可由下拉或圖表點選。",
        "重設篩選：回到 114 年、全部系所／教師／次項目。",
    ])
    add_h(doc, "3.2 指標與互動", 2)
    add_table(doc, ["項目", "判讀方式"], [
        ("總獎補助金額", "目前篩選資料的金額加總。"),
        ("教師數量", "目前資料中的不重複教師人數。"),
        ("年度占比", "目前篩選金額占所選年度基準金額的比例。"),
        ("教師平均獎補助", "總金額除以不重複教師人數。"),
        ("資料筆數", "目前符合條件的原始記錄筆數。"),
        ("次項目圖／年度趨勢", "點選次項目可交叉篩選；按清除選項還原。"),
        ("教師金額分布", "按顯示後查看各教師金額與比例。"),
    ], [3000, 6360])

    add_h(doc, "4. 教學精進")
    add_p(doc, "用途：分析教學獎勵主類、次類、細項、系所、教師參與、共同申請及金額／點數。")
    add_h(doc, "4.1 篩選與階層選取", 2)
    add_bullets(doc, [
        "可依年度、系所、教師、獎勵主類與獎勵次類篩選。",
        "主類、次類、細項圖表具有階層關係；上層選取會限制下層可見資料。",
        "圖表右上「清除選項」只清除對應圖表層級；「重設篩選」清除全部篩選與圖表選取。",
    ])
    add_h(doc, "4.2 指標與詳細資料", 2)
    add_table(doc, ["項目", "說明"], [
        ("總點數／總金額", "目前條件下的點數與金額加總。"),
        ("申請件數", "不重複申請編號數。"),
        ("共同申請件數", "申請型態為共同的不重複申請數。"),
        ("申請教師數", "不重複教師數。"),
        ("共同申請合作關係", "顯示共同申請所形成的教師合作連結。"),
        ("具體事項", "按「顯示具體事項」查看年度、教師、系所、成果、點數與金額。"),
        ("教師金額分布", "按鈕展開後查看目前條件下各教師的金額。"),
    ], [3000, 6360])
    add_note(doc, "隱私", "此頁教師姓名及共同申請名單採 O 字遮罩。")

    add_h(doc, "5. 論文發表")
    add_p(doc, "用途：以去重後論文為單位，分析收錄分類、作者角色、校內合著、出版國家與年度趨勢。")
    add_h(doc, "5.1 篩選", 2)
    add_bullets(doc, [
        "年度、系所、教師為單選；教師欄可搜尋。",
        "收錄分類為多選，符合任一選取分類的論文會保留。",
        "按「重設篩選」可回到全部資料。",
    ])
    add_h(doc, "5.2 重要指標定義", 2)
    add_table(doc, ["指標", "定義"], [
        ("去重後論文篇數", "以發表年份與正規化論文名稱整併相同論文。"),
        ("SSCI only", "含 SSCI、不含 SCIE。"),
        ("SCIE only", "含 SCIE、不含 SSCI；資料中的 SCI 會正規化為 SCIE。"),
        ("SSCI + SCIE", "同時含 SSCI 與 SCIE。"),
        ("第一作者篇數", "至少一位校內教師標記為第一作者。"),
        ("通訊作者篇數", "至少一位校內教師標記為通訊作者。"),
        ("第一/通訊作者篇數", "同一篇中有教師同時屬第一與通訊作者。"),
        ("校內合著篇數", "同一篇去重論文含兩位以上校內教師。"),
    ], [3300, 6060])
    add_p(doc, "圖表包含收錄分類、近三年趨勢、出版國家與校內合著關聯；需要教師發表排行時按「顯示教師排行」。")

    add_h(doc, "6. 計畫承接")
    add_p(doc, "用途：追蹤計畫件數、總金額、年度趨勢、計畫類型與系所／教師承接排行。")
    add_h(doc, "6.1 篩選與圖表", 2)
    add_bullets(doc, [
        "年度、系所、教師為單選；專案類型為多選。",
        "計畫件數以不重複識別號計算，總金額為符合條件記錄的計畫總金額加總。",
        "可點選年度趨勢、專案類型、系所排行與教師排行作為交叉選取。",
        "圖表選取可用「清除選項」移除；「重設篩選」會一併還原篩選與圖表選取。",
    ])
    add_h(doc, "6.2 計畫清單", 2)
    add_steps(doc, [
        "先套用適當的年度、系所、教師或計畫類型，縮小資料範圍。",
        "按「顯示計畫清單」。",
        "查看計畫名稱、案號、類型、期間、金額及國內委託單位。",
        "完成後按「收合計畫清單」。",
    ])

    add_h(doc, "7. 教師聘任")
    add_p(doc, "用途：檢視教師人力結構、職級與系所分布，並將教師主檔與論文、計畫、專利、技術移轉資料整合為研究評鑑風險提示。")
    add_h(doc, "7.1 篩選與摘要", 2)
    add_bullets(doc, [
        "篩選：系所、專兼任、編制、職級（可多選）。",
        "摘要：教師總數、專任教師、兼任教師、編制內與平均年齡。",
        "點選系所教師數圖或職級分布圖可交叉聚焦；按清除選項還原。",
        "按「顯示教師清單」查看個別聘任、學歷與學術專長。",
    ])
    add_h(doc, "7.2 研究評鑑風險名單", 2)
    add_p(doc, "按「顯示風險名單」後，系統列出符合候選條件且缺少計畫或研究成果的教師。候選條件為專任、講師以上、服務滿三年且未滿 62 歲。")
    add_table(doc, ["等級", "判定"], [
        ("高風險", "計畫件數為 0，且論文＋專利＋技術移轉也為 0。"),
        ("中風險", "僅缺少計畫或僅缺少研究成果其中一項。"),
    ], [2400, 6960])
    add_note(doc, "判讀限制", "風險名單是資料條件提示，不是正式考核結果。姓名差異、資料期間、漏登或 CSV 更新延遲都可能造成誤判；使用前應回查原始資料與適用規章。", "gold")

    add_h(doc, "8. 建議分析流程")
    add_steps(doc, [
        "先確認頁面標題、資料期間與用途是否符合問題。",
        "由年度開始，再逐步加入系所、教師與分類條件。",
        "記錄目前篩選條件後再閱讀摘要指標。",
        "使用圖表點選找出異常、集中或趨勢，再展開詳細清單查證。",
        "必要時切換其他模組交叉比對，例如教師頁的風險提示搭配論文與計畫頁查核。",
        "正式引用數字前，確認資料更新日期並保留條件截圖或人工紀錄。",
    ])

    add_h(doc, "9. 常見問題排除")
    add_table(doc, ["現象", "處理方式"], [
        ("畫面沒有資料", "按重設篩選；確認是否同時套用了過多條件；若仍無資料，請維運人員確認 CSV。"),
        ("教師搜尋不到", "先清除系所／年度等條件，再輸入姓名；名單可能會隨其他條件縮小。"),
        ("圖表數字與原始 CSV 列數不同", "確認是否套用篩選；論文會去重，計畫與申請件數也可能以識別號計算。"),
        ("點選圖表後其他區塊改變", "這是交叉篩選；按對應的清除選項或重設篩選。"),
        ("剛更新 CSV 但畫面沒變", "確認檔名映射、部署版本與服務是否已重新建置／啟動；再強制重新整理瀏覽器。"),
        ("姓名顯示 O", "獎補助與教學頁採個資遮罩，屬預期行為。"),
        ("風險名單似乎不正確", "核對教師姓名是否能跨檔匹配、服務年資／年齡與資料期間，並回查論文、計畫、專利及技轉來源。"),
    ], [3300, 6060])

    add_h(doc, "10. 資料與隱私注意事項")
    add_bullets(doc, [
        "僅在授權環境使用，不要將含教師個資的畫面或檔案任意分享。",
        "教師聘任與風險名單會顯示完整姓名及人事／研究資訊，截圖前應確認收件人與用途。",
        "畫面為分析輔助；人事、獎補助或評鑑決策應依正式規章與原始資料複核。",
        "系統目前沒有登入與操作稽核，離開座位時請鎖定裝置並遵循校內資訊安全政策。",
    ])

    add_h(doc, "附錄 A｜快速操作對照")
    add_table(doc, ["需求", "建議入口與操作"], [
        ("查某年度某系所獎補助", "獎補助分析 → 年度 → 系所 → 查看總額與次項目。"),
        ("查教學獎勵細項", "教學精進 → 年度／系所 → 點主類、次類、細項 → 顯示具體事項。"),
        ("查 SSCI／SCIE 論文", "論文發表 → 收錄分類多選 → 查看互斥桶與趨勢。"),
        ("查教師計畫", "計畫承接 → 教師搜尋 → 顯示計畫清單。"),
        ("查人力結構", "教師聘任 → 系所／專兼任／編制／職級 → 查看摘要與分布。"),
        ("查研究評鑑提醒", "教師聘任 → 套用系所等條件 → 顯示風險名單 → 回查來源資料。"),
    ], [3600, 5760])
    path = OUT / "使用者操作手冊.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for output in (build_architecture(), build_manual()):
        print(output)
