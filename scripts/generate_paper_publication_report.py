from __future__ import annotations

import csv
import math
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "papers_112_114.csv"
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "paper_publication_assets"
DOCX_PATH = OUT_DIR / "期刊論文發表情形分析報告.docx"

FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msjh.ttc"),
    Path("C:/Windows/Fonts/msjhbd.ttc"),
    Path("C:/Windows/Fonts/mingliu.ttc"),
    Path("C:/Windows/Fonts/NotoSansCJK-Regular.ttc"),
]

ACCENT = "1f6f86"
ACCENT_2 = "0f766e"
TEXT = "1f2937"
MUTED = "64748b"
GRID = "dbe4ea"
PALETTE = ["14b8a6", "2563eb", "f59e0b", "db2777", "16a34a", "7c3aed", "ef4444", "0891b2"]


@dataclass(frozen=True)
class Record:
    school_year: str
    department: str
    teacher: str
    appointment: str
    title: str
    journal: str
    categories: tuple[str, ...]
    author_order: str
    is_corresponding: bool
    has_peer_review: bool
    is_international: bool
    country: str
    publication_year: int
    publication_format: str


@dataclass(frozen=True)
class Publication:
    school_year: str
    publication_year: int
    title: str
    journal: str
    teachers: tuple[str, ...]
    departments: tuple[str, ...]
    categories: tuple[str, ...]
    corresponding_authors: tuple[str, ...]
    first_authors: tuple[str, ...]
    has_peer_review: bool
    is_international: bool
    has_internal_coauthor: bool
    countries: tuple[str, ...]
    publication_format: str


def clean(value: str | None) -> str:
    return re.sub(r"\s+", " ", (value or "").replace("\u00a0", " ")).strip()


def normalize_title(value: str) -> str:
    return re.sub(r"[^\w]+", "", clean(value).lower(), flags=re.UNICODE)


def split_categories(value: str) -> tuple[str, ...]:
    items = []
    for item in clean(value).split(","):
        item = clean(item)
        if item == "SCI":
            item = "SCIE"
        if item:
            items.append(item)
    return tuple(items)


def unique_sorted(values: list[str]) -> tuple[str, ...]:
    return tuple(sorted(set(v for v in values if v), key=lambda item: item))


def pct(part: int, whole: int) -> float:
    return round(part * 100 / whole, 1) if whole else 0.0


def read_publications() -> tuple[int, list[Record], list[Publication]]:
    with DATA_PATH.open("r", encoding="utf-8-sig", newline="") as file:
        rows = list(csv.DictReader(file))

    records: list[Record] = []
    for row in rows:
        try:
            publication_year = int(re.sub(r"[^0-9-]", "", clean(row.get("發表年份"))))
        except ValueError:
            publication_year = 0

        record = Record(
            school_year=clean(row.get("年度")),
            department=clean(row.get("主聘系所")),
            teacher=clean(row.get("教師姓名")),
            appointment=clean(row.get("專兼任")),
            title=clean(row.get("論文名稱")),
            journal=clean(row.get("刊物名稱")),
            categories=split_categories(clean(row.get("論文收錄分類"))),
            author_order=clean(row.get("作者順序")),
            is_corresponding=clean(row.get("通訊作者")) == "是",
            has_peer_review=clean(row.get("是否具有審稿制度")) == "是",
            is_international=clean(row.get("期刊論文是否為跨國(地區)合作")) == "是",
            country=clean(row.get("期刊出版地國家/地區")),
            publication_year=publication_year,
            publication_format=clean(row.get("發表型式")),
        )
        if (
            record.appointment == "專任"
            and record.school_year
            and record.department
            and record.teacher
            and record.title
            and record.publication_year > 0
        ):
            records.append(record)

    grouped: dict[str, list[Record]] = defaultdict(list)
    for record in records:
        grouped[f"{record.publication_year}::{normalize_title(record.title)}"].append(record)

    publications: list[Publication] = []
    for group in grouped.values():
        base = group[0]
        teachers = unique_sorted([item.teacher for item in group])
        publications.append(
            Publication(
                school_year=base.school_year,
                publication_year=base.publication_year,
                title=base.title,
                journal=base.journal,
                teachers=teachers,
                departments=unique_sorted([item.department for item in group]),
                categories=unique_sorted([cat for item in group for cat in item.categories]),
                corresponding_authors=unique_sorted(
                    [item.teacher for item in group if item.is_corresponding]
                ),
                first_authors=unique_sorted(
                    [item.teacher for item in group if item.author_order == "第一作者"]
                ),
                has_peer_review=any(item.has_peer_review for item in group),
                is_international=any(item.is_international for item in group),
                has_internal_coauthor=len(teachers) > 1,
                countries=unique_sorted([item.country for item in group]),
                publication_format=base.publication_format,
            )
        )

    publications.sort(key=lambda item: (item.school_year, item.title))
    return len(rows), records, publications


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [Path("C:/Windows/Fonts/msjhbd.ttc")] if bold else []
    candidates.extend(FONT_CANDIDATES)
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def hex_color(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_text(draw: ImageDraw.ImageDraw, xy, text, font, fill=TEXT, anchor=None):
    draw.text(xy, text, font=font, fill=hex_color(fill), anchor=anchor)


def save_year_trend(year_counts: list[tuple[str, int]], path: Path) -> None:
    width, height = 1280, 640
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, True)
    label_font = load_font(24)
    small_font = load_font(21)
    draw_text(draw, (62, 44), "112-114 學年度期刊論文發表趨勢", title_font, ACCENT)

    left, top, right, bottom = 110, 130, 1190, 535
    max_value = max(value for _, value in year_counts)
    y_max = math.ceil(max_value / 50) * 50
    for i in range(6):
        value = y_max * i / 5
        y = bottom - (bottom - top) * i / 5
        draw.line((left, y, right, y), fill=hex_color(GRID), width=2)
        draw_text(draw, (left - 18, y), f"{int(value)}", small_font, MUTED, "rm")

    points = []
    step = (right - left) / (len(year_counts) - 1)
    for index, (year, value) in enumerate(year_counts):
        x = left + index * step
        y = bottom - (bottom - top) * value / y_max
        points.append((x, y))
        draw.rounded_rectangle((x - 52, y, x + 52, bottom), radius=18, fill=hex_color("dbeafe"))
        draw.rounded_rectangle((x - 52, y, x + 52, bottom), radius=18, outline=hex_color("bfdbfe"), width=2)
        draw.ellipse((x - 11, y - 11, x + 11, y + 11), fill=hex_color("2563eb"))
        draw_text(draw, (x, y - 38), str(value), label_font, "0f172a", "mm")
        draw_text(draw, (x, bottom + 36), year, label_font, TEXT, "mm")

    draw.line(points, fill=hex_color("14b8a6"), width=8, joint="curve")
    image.save(path)


def save_horizontal_bar(title: str, data: list[tuple[str, int]], path: Path, color: str = "14b8a6") -> None:
    width, height = 1280, 760
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, True)
    label_font = load_font(23)
    value_font = load_font(22, True)
    draw_text(draw, (62, 42), title, title_font, ACCENT)

    left, top, right = 360, 112, 1130
    row_h = 58
    max_value = max(value for _, value in data) if data else 1
    for index, (name, value) in enumerate(data):
        y = top + index * row_h
        draw_text(draw, (left - 20, y + 18), name, label_font, TEXT, "rm")
        bar_w = (right - left) * value / max_value
        draw.rounded_rectangle((left, y, right, y + 34), radius=17, fill=hex_color("eef2f7"))
        draw.rounded_rectangle((left, y, left + bar_w, y + 34), radius=17, fill=hex_color(color))
        draw_text(draw, (left + bar_w + 16, y + 17), str(value), value_font, TEXT, "lm")
    image.save(path)


def save_category_chart(data: list[tuple[str, int]], path: Path) -> None:
    width, height = 1280, 640
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, True)
    label_font = load_font(24)
    small_font = load_font(22)
    draw_text(draw, (62, 44), "論文收錄分類分布", title_font, ACCENT)

    total = sum(value for _, value in data)
    box = (90, 135, 520, 565)
    start = -90
    for index, (_, value) in enumerate(data):
        extent = 360 * value / total
        draw.pieslice(box, start=start, end=start + extent, fill=hex_color(PALETTE[index % len(PALETTE)]))
        start += extent
    draw.ellipse((205, 250, 405, 450), fill="#ffffff")
    draw_text(draw, (305, 315), f"{total}", load_font(34, True), TEXT, "mm")
    draw_text(draw, (305, 358), "篇次", small_font, MUTED, "mm")

    x0, y0 = 620, 150
    for index, (name, value) in enumerate(data):
        y = y0 + index * 72
        color = PALETTE[index % len(PALETTE)]
        draw.rounded_rectangle((x0, y, x0 + 34, y + 34), radius=8, fill=hex_color(color))
        draw_text(draw, (x0 + 52, y + 5), name, label_font, TEXT)
        draw_text(draw, (x0 + 300, y + 5), f"{value} 篇次", label_font, TEXT)
        draw_text(draw, (x0 + 480, y + 7), f"{pct(value, total)}%", small_font, MUTED)
    image.save(path)


def save_role_chart(metrics: list[tuple[str, int, float]], path: Path) -> None:
    width, height = 1280, 560
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, True)
    label_font = load_font(24)
    value_font = load_font(34, True)
    draw_text(draw, (62, 44), "教師作者角色與合作型態", title_font, ACCENT)
    card_w, gap, top = 260, 30, 150
    for index, (name, value, ratio) in enumerate(metrics):
        x = 70 + index * (card_w + gap)
        color = PALETTE[index % len(PALETTE)]
        draw.rounded_rectangle((x, top, x + card_w, top + 270), radius=28, fill=hex_color("f8fafc"), outline=hex_color("dbe4ea"), width=2)
        draw.rounded_rectangle((x, top, x + card_w, top + 16), radius=8, fill=hex_color(color))
        draw_text(draw, (x + card_w / 2, top + 80), f"{value}", value_font, color, "mm")
        draw_text(draw, (x + card_w / 2, top + 130), f"{ratio:.1f}%", load_font(28, True), TEXT, "mm")
        lines = name.split("/")
        for line_index, line in enumerate(lines):
            draw_text(draw, (x + card_w / 2, top + 190 + line_index * 34), line, label_font, MUTED, "mm")
    image.save(path)


def counter_by(publications: list[Publication], getter) -> Counter:
    counter: Counter = Counter()
    for publication in publications:
        values = getter(publication)
        if isinstance(values, str):
            values = [values]
        for value in values:
            if value:
                counter[value] += 1
    return counter


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill: str = "e0f2fe") -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "110")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.color.rgb = RGBColor.from_string(ACCENT if level == 1 else ACCENT_2)
    return paragraph


def add_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TEXT)
    return paragraph


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, width_inches: float = 6.25):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.font.name = "Microsoft JhengHei"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
    style_table(table)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def build_report() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    raw_row_count, records, publications = read_publications()

    year_counts = sorted(counter_by(publications, lambda p: p.school_year).items(), key=lambda item: int(item[0]))
    category_counts = counter_by(publications, lambda p: p.categories or ("未分類",)).most_common()
    department_counts = counter_by(publications, lambda p: p.departments).most_common(10)
    teacher_counts = counter_by(publications, lambda p: p.teachers).most_common(10)
    journal_counts = counter_by(publications, lambda p: p.journal).most_common(10)

    total = len(publications)
    peer_count = sum(1 for item in publications if item.has_peer_review)
    first_count = sum(1 for item in publications if item.first_authors)
    corr_count = sum(1 for item in publications if item.corresponding_authors)
    first_or_corr_count = sum(1 for item in publications if item.first_authors or item.corresponding_authors)
    internal_count = sum(1 for item in publications if item.has_internal_coauthor)
    intl_count = sum(1 for item in publications if item.is_international)
    indexed_count = sum(1 for item in publications if {"SCIE", "SSCI"} & set(item.categories))
    scie_only = sum(1 for item in publications if "SCIE" in item.categories and "SSCI" not in item.categories)
    ssci_only = sum(1 for item in publications if "SSCI" in item.categories and "SCIE" not in item.categories)
    both_ssci_scie = sum(1 for item in publications if "SSCI" in item.categories and "SCIE" in item.categories)

    pair_counter: Counter = Counter()
    for publication in publications:
        teachers = list(publication.teachers)
        for index, teacher in enumerate(teachers):
            for coauthor in teachers[index + 1 :]:
                pair_counter[f"{teacher} × {coauthor}"] += 1

    year_chart = ASSET_DIR / "year_trend.png"
    category_chart = ASSET_DIR / "category_distribution.png"
    department_chart = ASSET_DIR / "department_top10.png"
    role_chart = ASSET_DIR / "role_metrics.png"
    save_year_trend(year_counts, year_chart)
    save_category_chart(category_counts, category_chart)
    save_horizontal_bar("系所期刊論文發表量 Top 10", department_counts, department_chart)
    save_role_chart(
        [
            ("第一作者", first_count, pct(first_count, total)),
            ("通訊作者", corr_count, pct(corr_count, total)),
            ("第一或/通訊作者", first_or_corr_count, pct(first_or_corr_count, total)),
            ("校內合著", internal_count, pct(internal_count, total)),
        ],
        role_chart,
    )

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("期刊論文發表情形分析報告")
    run.bold = True
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("112 至 114 學年度 | 專任教師期刊論文資料")
    sub_run.font.name = "Microsoft JhengHei"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string(MUTED)

    add_heading(doc, "一、資料範圍與統計口徑")
    add_paragraph(
        doc,
        f"本報告依據 {DATA_PATH.name} 之 112 至 114 學年度期刊論文資料彙整，統計口徑採專任教師為主，並以「發表年份＋論文名稱」去除同篇論文重複列後分析。原始資料共 {raw_row_count:,} 筆，其中專任教師有效資料 {len(records):,} 筆，去重後共計 {total:,} 篇期刊論文，涵蓋 {len(set(r.teacher for r in records))} 位教師、{len(set(r.department for r in records))} 個系所單位。",
    )

    add_simple_table(
        doc,
        ["指標", "數值"],
        [
            ["專任教師有效資料列", f"{len(records):,} 筆"],
            ["去重後期刊論文", f"{total:,} 篇"],
            ["涵蓋教師", f"{len(set(r.teacher for r in records))} 位"],
            ["涵蓋系所", f"{len(set(r.department for r in records))} 個"],
            ["具審稿制度論文", f"{peer_count:,} 篇（{pct(peer_count, total):.1f}%）"],
        ],
        [7.5, 7.0],
    )

    add_heading(doc, "二、整體發表趨勢")
    add_paragraph(
        doc,
        "112 至 114 學年度期刊論文發表量整體維持穩定，114 學年度較 113 學年度增加 34 篇，顯示近期研究發表動能有回升跡象。113 學年度雖略為下降，但三年趨勢未呈現明顯衰退。",
    )
    add_picture_with_caption(doc, year_chart, "圖 1　112-114 學年度期刊論文發表趨勢")
    add_simple_table(doc, ["學年度", "論文篇數"], [[year, f"{value:,}"] for year, value in year_counts], [6, 5])

    add_heading(doc, "三、論文收錄與品質表現")
    add_paragraph(
        doc,
        f"論文收錄分類以 SCIE 為最大宗，其次為其他與 SSCI。若以 SCIE 或 SSCI 作為主要國際索引觀察，至少具 SCIE 或 SSCI 收錄者共 {indexed_count:,} 篇，占 {pct(indexed_count, total):.1f}%。其中 SCIE only {scie_only:,} 篇、SSCI only {ssci_only:,} 篇、同時具 SCIE 與 SSCI 者 {both_ssci_scie:,} 篇。",
    )
    add_picture_with_caption(doc, category_chart, "圖 2　論文收錄分類分布")
    add_simple_table(
        doc,
        ["收錄分類", "篇次"],
        [[name, f"{value:,}"] for name, value in category_counts],
        [7.0, 4.0],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "四、教師作者角色與合作型態")
    add_paragraph(
        doc,
        f"教師在多數論文中具有研究主導角色。第一作者論文 {first_count:,} 篇，占 {pct(first_count, total):.1f}%；通訊作者論文 {corr_count:,} 篇，占 {pct(corr_count, total):.1f}%；第一或通訊作者論文 {first_or_corr_count:,} 篇，占 {pct(first_or_corr_count, total):.1f}%。校內合著論文 {internal_count:,} 篇，占 {pct(internal_count, total):.1f}%，可作為後續推動跨系所研究社群與共同計畫申請的重要基礎。",
    )
    add_picture_with_caption(doc, role_chart, "圖 3　教師作者角色與校內合著比例")
    add_simple_table(
        doc,
        ["指標", "篇數", "占比"],
        [
            ["第一作者論文", f"{first_count:,}", f"{pct(first_count, total):.1f}%"],
            ["通訊作者論文", f"{corr_count:,}", f"{pct(corr_count, total):.1f}%"],
            ["第一或通訊作者論文", f"{first_or_corr_count:,}", f"{pct(first_or_corr_count, total):.1f}%"],
            ["校內合著論文", f"{internal_count:,}", f"{pct(internal_count, total):.1f}%"],
            ["跨國合作論文", f"{intl_count:,}", f"{pct(intl_count, total):.1f}%"],
        ],
        [7, 4, 4],
    )

    add_heading(doc, "五、系所與教師發表分布")
    add_paragraph(
        doc,
        "期刊論文主要集中於護理、健康照護與健康產業相關領域。護理系為最大發表單位，若加計嘉義分部護理系，兩者合計已占整體論文量超過半數，顯示學校研究能量與特色領域高度一致；同時也代表其他系所仍有擴大研究產出的空間。",
    )
    add_picture_with_caption(doc, department_chart, "圖 4　系所期刊論文發表量 Top 10")
    add_simple_table(
        doc,
        ["系所", "篇數"],
        [[name, f"{value:,}"] for name, value in department_counts],
        [9.0, 3.5],
    )

    add_heading(doc, "六、重點教師與期刊來源")
    add_paragraph(
        doc,
        "部分教師具高度且穩定的研究產出，適合作為研究社群、跨域合作或新進教師研究輔導的核心成員。期刊來源方面，除校內及護理專業期刊外，也包含 Nurse Education Today、Healthcare、BMC Nursing、Geriatric Nursing、Journal of Nursing Research 等國際期刊。",
    )
    add_simple_table(
        doc,
        ["教師", "篇數"],
        [[name, f"{value:,}"] for name, value in teacher_counts],
        [7.5, 4.0],
    )
    add_simple_table(
        doc,
        ["期刊名稱", "篇數"],
        [[name, f"{value:,}"] for name, value in journal_counts],
        [11.0, 3.0],
    )

    doc.add_page_break()
    add_heading(doc, "七、綜合研析與建議")
    for text in [
        "整體而言，112 至 114 學年度期刊論文發表量穩定，114 學年度回升，且 SCIE / SSCI 收錄比例良好，教師擔任第一或通訊作者比例高，顯示研究主導性明確。",
        "建議持續支持高產出教師與核心研究團隊，擴大其帶動效果，並以成熟合作網絡作為高品質期刊投稿與外部研究計畫申請的基礎。",
        "建議鼓勵非主要發表系所建立穩定研究主題，透過跨域共作、研究方法工作坊與校內合著媒合，降低研究能量過度集中情形。",
        "資料中跨國合作欄位標示為 0 篇，但期刊出版地涵蓋中華民國、美國、英國、瑞士等多個地區；後續可進一步檢核填報標準，區分「國外期刊出版」與「跨國合著」兩種不同概念。",
    ]:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        paragraph.paragraph_format.space_after = Pt(7)
        run = paragraph.add_run("• ")
        run.font.color.rgb = RGBColor.from_string(ACCENT_2)
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        body = paragraph.add_run(text)
        body.font.name = "Microsoft JhengHei"
        body._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        body.font.size = Pt(11)
        body.font.color.rgb = RGBColor.from_string(TEXT)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("期刊論文發表情形分析報告")
    footer_run.font.name = "Microsoft JhengHei"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    core_properties = doc.core_properties
    core_properties.title = "期刊論文發表情形分析報告"
    core_properties.author = "Codex"
    core_properties.last_modified_by = "Codex"
    doc.save(DOCX_PATH)

    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
